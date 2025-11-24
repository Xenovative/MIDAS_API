"""
Document parser for various file formats
Supports: PDF, DOC, DOCX, TXT, JSON
"""
import json
from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document as DocxDocument


class DocumentParser:
    """Parse various document formats to text"""
    
    @staticmethod
    def parse_file(file_path: Path, filename: str) -> str:
        """
        Parse a file and extract text content
        
        Args:
            file_path: Path to the file
            filename: Original filename (for extension detection)
            
        Returns:
            Extracted text content
        """
        extension = filename.lower().split('.')[-1]
        
        if extension == 'pdf':
            return DocumentParser._parse_pdf(file_path)
        elif extension in ['doc', 'docx']:
            return DocumentParser._parse_docx(file_path)
        elif extension == 'json':
            return DocumentParser._parse_json(file_path)
        elif extension == 'txt':
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    @staticmethod
    def parse_bytes(content: bytes, filename: str) -> str:
        """
        Parse file content from bytes
        
        Args:
            content: File content as bytes
            filename: Original filename (for extension detection)
            
        Returns:
            Extracted text content
        """
        import tempfile
        
        # Write to temp file and parse
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{filename.split('.')[-1]}") as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)
        
        try:
            return DocumentParser.parse_file(tmp_path, filename)
        finally:
            tmp_path.unlink()
    
    @staticmethod
    def _parse_pdf(file_path: Path) -> str:
        """Extract text from PDF"""
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                        continue
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
        
        if not text_parts:
            raise ValueError("No text content could be extracted from PDF")
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _parse_docx(file_path: Path) -> str:
        """Extract text from DOCX (also works for DOC if converted)"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            if not text_parts:
                raise ValueError("No text content found in document")
            
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def _parse_json(file_path: Path) -> str:
        """Parse JSON and convert to readable text"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to formatted text
            if isinstance(data, dict):
                return DocumentParser._dict_to_text(data)
            elif isinstance(data, list):
                return DocumentParser._list_to_text(data)
            else:
                return json.dumps(data, indent=2)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON format: {str(e)}")
        except Exception as e:
            raise ValueError(f"Failed to parse JSON: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_path: Path) -> str:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                raise ValueError(f"Failed to read text file: {str(e)}")
    
    @staticmethod
    def _dict_to_text(data: dict, indent: int = 0) -> str:
        """Convert dictionary to readable text"""
        lines = []
        prefix = "  " * indent
        
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}{key}:")
                lines.append(DocumentParser._dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{prefix}{key}:")
                lines.append(DocumentParser._list_to_text(value, indent + 1))
            else:
                lines.append(f"{prefix}{key}: {value}")
        
        return "\n".join(lines)
    
    @staticmethod
    def _list_to_text(data: list, indent: int = 0) -> str:
        """Convert list to readable text"""
        lines = []
        prefix = "  " * indent
        
        for i, item in enumerate(data):
            if isinstance(item, dict):
                lines.append(f"{prefix}[{i}]:")
                lines.append(DocumentParser._dict_to_text(item, indent + 1))
            elif isinstance(item, list):
                lines.append(f"{prefix}[{i}]:")
                lines.append(DocumentParser._list_to_text(item, indent + 1))
            else:
                lines.append(f"{prefix}[{i}]: {item}")
        
        return "\n".join(lines)
    
    @staticmethod
    def get_supported_extensions() -> list[str]:
        """Get list of supported file extensions"""
        return ['txt', 'pdf', 'doc', 'docx', 'json']
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file format is supported"""
        extension = filename.lower().split('.')[-1]
        return extension in DocumentParser.get_supported_extensions()


# Global parser instance
document_parser = DocumentParser()
